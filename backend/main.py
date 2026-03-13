import os
import sys
import ee
import json
import math
import logging
import io
import numpy as np
import requests
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import concurrent.futures
from statsmodels.tsa.seasonal import STL
import ruptures as rpt
from scipy import stats

from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dateutil.relativedelta import relativedelta

from fastapi import FastAPI, HTTPException, Header, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import geopandas as gpd
import tempfile
import shutil
import zipfile
import fiona
import pandas as pd
from shapely.ops import transform

def to_2d(x, y, z=None):
    """Force 2D coordinates for GE compatibility."""
    return (x, y)

fiona.drvsupport.supported_drivers['KML'] = 'rw' # Enable KML support
fiona.drvsupport.supported_drivers['LIBKML'] = 'rw'

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

# ==========================================
# MODULE: UTILS & LOGGING
# ==========================================

def setup_logging(log_file: str = 'wetland_analysis.log', level: int = logging.INFO) -> logging.Logger:
    """Configure logging for the application."""
    logger = logging.getLogger('wetland_monitor')
    logger.setLevel(level)
    if logger.handlers:
        return logger
    
    detailed_formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    simple_formatter = logging.Formatter('%(message)s')
    
    # Check if we can write to log file, otherwise skip file logging
    try:
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(detailed_formatter)
        logger.addHandler(file_handler)
    except Exception:
        pass # Skip file logging if permission denied
    
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(simple_formatter)
    logger.addHandler(console_handler)
    
    return logger

def format_analysis_summary(stats: dict, mode: str) -> str:
    """Format analysis statistics into human-readable string."""
    if not stats:
        return f"{mode}: No data available"
    return f"""
{mode} Analysis Summary:
  Current Median: {stats.get('median', 'N/A'):.4f}
  Range: [{stats.get('min', 'N/A'):.4f}, {stats.get('max', 'N/A'):.4f}]
  Std Dev: {stats.get('std', 'N/A'):.4f}
  Data Points: {stats.get('count', 0)}
  CV: {stats.get('cv', 'N/A'):.2f}%
""".strip()

def create_error_response(error: Exception, mode: str = None) -> dict:
    """Create standardized error response."""
    return {
        "status": "error",
        "error": str(error),
        "error_type": type(error).__name__,
        "mode": mode,
        "timestamp": datetime.now().isoformat()
    }

def log_process_stage(stage: str, mode: str = None, status: str = 'processing') -> str:
    """Create formatted log message for process stages."""
    icons = {'processing': '⚙️', 'completed': '✓', 'error': '✗', 'info': 'ℹ️'}
    icon = icons.get(status, '')
    timestamp = datetime.now().strftime("%H:%M:%S")
    
    if mode:
        msg = f"[{timestamp}] {icon}  {mode}"
        if stage: msg += f": {stage}"
    else:
        msg = f"[{timestamp}] {icon}  {stage}"
    
    if status == 'completed' and not stage:
        msg += " completed"
        
    return msg

# Initialize Logger
logger = setup_logging()

# ==========================================
# MODULE: VALIDATORS
# ==========================================

class ValidationError(Exception):
    """Custom exception for validation failures"""
    pass

def validate_date_range(start_date_str: str, end_date_str: str, max_range_days: int = 4015) -> Tuple[datetime, datetime]:
    """Validate date range inputs."""
    try:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
    except ValueError as e:
        raise ValidationError(f"Invalid date format. Use YYYY-MM-DD. Error: {e}")
    
    if start_date >= end_date:
        raise ValidationError("start_date must be before end_date")
    
    if (end_date - start_date).days < 7:
        raise ValidationError("Date range must be at least 7 days")
        
    return start_date, end_date

def validate_geometry(geojson: Dict[str, Any], min_area_km2: float = 0.01, max_area_km2: float = 1000) -> ee.Geometry:
    """Validate GeoJSON geometry for analysis."""
    try:
        geometry = geojson.get('geometry')
        if not geometry:
            logger.error("Missing 'geometry' field in GeoJSON")
            raise ValidationError("Missing 'geometry' field in GeoJSON")
        
        geom_type = geometry.get('type')
        
        # Force 2D geometries (GEE can fail with 3D from KML)
        # Simple recursion to strip Z if present in coordinates
        def strip_z(coords):
            if not isinstance(coords, (list, tuple)):
                return coords
            if len(coords) > 0 and isinstance(coords[0], (int, float)):
                return list(coords[:2])
            return [strip_z(c) for c in coords]
        
        if 'coordinates' in geometry:
            geometry['coordinates'] = strip_z(geometry['coordinates'])

        aoi = ee.Geometry(geometry)
        aoi = aoi.simplify(maxError=1)
        return aoi
    except Exception as e:
        logger.error(f"Geometry validation failed: {e}")
        raise ValidationError(f"Geometry validation failed: {e}")

# ==========================================
# MODULE: ROBUST STATS
# ==========================================

def calculate_robust_statistics(data: List[Dict[str, Any]]) -> Optional[Dict[str, float]]:
    """Calculate robust statistics resistant to outliers."""
    values = [d['value'] for d in data if d.get('value') is not None]
    if len(values) < 1: return None
    
    values_array = np.array(values)
    mean_val = np.mean(values_array)
    std_val = np.std(values_array)
    
    # Robust CV: Use absolute mean to avoid negative CVs
    # If mean is extremely close to zero, CV becomes less meaningful
    abs_mean = abs(mean_val)
    cv = (std_val / abs_mean * 100) if abs_mean > 1e-6 else 0
    p25 = np.percentile(values_array, 25)
    p75 = np.percentile(values_array, 75)
    
    return {
        'mean': float(mean_val),
        'median': float(np.median(values_array)),
        'std': float(std_val),
        'min': float(np.min(values_array)),
        'max': float(np.max(values_array)),
        'count': len(values),
        'cv': float(cv),
        'iqr': float(p75 - p25)
    }

def detect_outliers(data: List[Dict[str, Any]], method: str = 'iqr', threshold: float = 1.5) -> List[Dict[str, Any]]:
    """Detect and flag outliers in time series data."""
    values = [d['value'] for d in data if d.get('value') is not None]
    if len(values) < 4:
        for d in data: d['is_outlier'] = False
        return data
    
    values_array = np.array(values)
    q1 = np.percentile(values_array, 25)
    q3 = np.percentile(values_array, 75)
    iqr = q3 - q1
    lower = q1 - threshold * iqr
    upper = q3 + threshold * iqr
    
    for d in data:
        if d.get('value') is not None:
            d['is_outlier'] = bool(d['value'] < lower or d['value'] > upper)
        else:
            d['is_outlier'] = False
    return data

def validate_temporal_coverage(data: List[Dict[str, Any]], min_days: int = 30) -> Dict[str, Any]:
    """Validate that time series data has adequate temporal coverage."""
    if not data or len(data) < 2:
        return {'valid': False, 'reason': 'Insufficient data points', 'coverage_days': 0}
    
    dates = []
    for d in data:
        try: dates.append(datetime.strptime(d['date'], '%Y-%m-%d'))
        except: continue
        
    if len(dates) < 2:
        return {'valid': False, 'reason': 'Invalid dates', 'coverage_days': 0}
        
    coverage_days = (max(dates) - min(dates)).days
    return {
        'valid': coverage_days >= min_days,
        'reason': 'Adequate coverage' if coverage_days >= min_days else 'Insufficient coverage',
        'coverage_days': coverage_days,
        'data_points': len(data),
        'start_date': min(dates).strftime('%Y-%m-%d') if dates else 'N/A',
        'end_date': max(dates).strftime('%Y-%m-%d') if dates else 'N/A'
    }

def calculate_spatial_consistency(image, area, scale=10):
    """
    Calculate Global Moran's I for spatial autocorrelation on an Earth Engine image.
    Helps determine if changes are clustered (real) or random noise.
    """
    try:
        if not image: return 0.0
        
        # 1. Normalize the image (standardize)
        stats = image.reduceRegion(reducer=ee.Reducer.mean().combine(ee.Reducer.stdDev(), "", True), geometry=area, scale=scale)
        mean = ee.Number(stats.get('Value_mean'))
        std = ee.Number(stats.get('Value_stdDev'))
        
        # Avoid division by zero
        standardized = image.subtract(mean).divide(std.max(0.00001))
        
        # 2. Local Moran's I approximation via reduction
        # Define a 3x3 kernel for adjacency
        weights = [
            [1, 1, 1],
            [1, 0, 1],
            [1, 1, 1]
        ]
        kernel = ee.Kernel.fixed(3, 3, weights)
        
        # Calculate sum of neighbors
        neighbors_sum = standardized.reduceNeighborhood(reducer=ee.Reducer.sum(), kernel=kernel)
        
        # Moran's I = standardized * neighbors_sum / normalization
        moran_image = standardized.multiply(neighbors_sum)
        
        # Global Moran's I (mean over area)
        global_moran = moran_image.reduceRegion(reducer=ee.Reducer.mean(), geometry=area, scale=scale).get('Value')
        
        val = ee.Number(global_moran).getInfo()
        return float(val) if val is not None else 0.0
    except Exception as e:
        logger.warning(f"Spatial consistency calculation failed: {e}")
        return 0.0

def calculate_trend_statistics(current_data: List[Dict], previous_data: List[Dict]) -> Optional[Dict[str, float]]:
    """Calculate trend statistics comparing two periods with robust P-values."""
    current_stats = calculate_robust_statistics(current_data)
    previous_stats = calculate_robust_statistics(previous_data)
    
    if not current_stats: return None
    
    curr_values = [d['value'] for d in current_data if d.get('value') is not None]
    prev_values = [d['value'] for d in previous_data if d.get('value') is not None] if previous_data else []
    
    # Statistical significance via T-Test
    p_value = 1.0
    if len(curr_values) > 1 and len(prev_values) > 1:
        t_stat, p_val = stats.ttest_ind(curr_values, prev_values, equal_var=False)
        p_value = float(p_val)

    curr_med = current_stats['median']
    prev_med = previous_stats['median'] if previous_stats else None
    
    trend_pct = None
    if prev_med is not None:
        # Dampening: If previous median is near zero, use a small constant to avoid extreme percentages
        # Increase threshold to 0.02 for more stability in dry/unproductive areas
        denom = max(abs(prev_med), 0.02) 
        trend_pct = ((curr_med - prev_med) / denom) * 100
            
    if trend_pct is not None:
        trend_pct = max(min(trend_pct, 1000), -1000)
    
    return {
        'previous_median': prev_med,
        'current_median': curr_med,
        'trend_percent': trend_pct,
        'p_value': p_value,
        'is_significant': p_value < 0.05 if p_value is not None else False,
        'absolute_change': (curr_med - prev_med) if prev_med is not None else None
    }

def apply_bfast_analysis(time_series: List[Dict], mode: str) -> Optional[Dict[str, Any]]:
    """
    Apply a robust BFAST-like analysis to the time series.
    Detects trend, seasonality, and breakpoints (structural changes).
    """
    if not time_series or len(time_series) < 12: # Minimum a year for basic trend
        return None
    
    try:
        # 1. Prepare Data
        df = pd.DataFrame(time_series)
        df['date'] = pd.to_datetime(df['date'])
        df = df.sort_values('date').set_index('date')
        
        # Ensure 'value' exists and is numeric
        if 'value' not in df.columns: return None
        df['value'] = pd.to_numeric(df['value'], errors='coerce')
        # Resample to monthly to handle irregular GEE data
        valid_count = len(df)
        df_resampled_base = df['value'].resample('MS').mean()
        
        # If data is too sparse (< 30% of monthly slots filled), skip interpolation
        if valid_count / len(df_resampled_base) < 0.3:
            return None
            
        # Using PCHIP interpolation for smoother, shape-preserving environmental series
        df_resampled = df_resampled_base.interpolate(method='pchip')
        
        # BFAST needs at least 2 full periods for seasonal decomposition
        # If less, we skip seasonal and just do trend/break detection
        has_seasonality = len(df_resampled) >= 24
        
        if len(df_resampled) < 6: return None # Absolute minimum

        # 2. Decomposition
        if has_seasonality:
            # STL Decomposition (Seasonal-Trend decomposition using Loess)
            stl = STL(df_resampled, period=12, robust=True).fit()
            trend = stl.trend
            seasonal = stl.seasonal
            remainder = stl.resid
        else:
            # Fallback for short series: Simple rolling mean or just the original for trend
            trend = df_resampled.rolling(window=3, min_periods=1, center=True).mean()
            seasonal = pd.Series(0, index=df_resampled.index)
            remainder = df_resampled - trend

        # 3. Structural Change Detection (Breakpoints)
        # Using PELT (Pruned Exact Linear Time) on the Trend component
        # We use L2 (least squares) cost and a penalty based on data size
        signal = trend.values.reshape(-1, 1)
        # h (min_size) is the minimum segment length (6 months)
        algo = rpt.Pelt(model="l2", min_size=6).fit(signal)
        # Dynamic penalty: Lower factor for better sensitivity in wetland monitoring
        pen = np.log(len(signal)) * 1.5 if len(signal) > 10 else 1
        breakpoints_indices = algo.predict(pen=pen)
        
        # 4. Segment Analysis
        segments = []
        last_bk = 0
        for bk in breakpoints_indices:
            if bk > len(trend): bk = len(trend)
            if bk <= last_bk: continue
            
            seg_data = trend.iloc[last_bk:bk]
            if len(seg_data) >= 2:
                x = np.arange(len(seg_data))
                slope, intercept, r_value, p_value, std_err = stats.linregress(x, seg_data.values)
                
                segments.append({
                    "start_date": trend.index[last_bk].strftime("%Y-%m-%d"),
                    "end_date": trend.index[bk-1].strftime("%Y-%m-%d"),
                    "slope": float(slope),
                    "slope_label": "Incremento" if slope > 0.001 else "Descenso" if slope < -0.001 else "Estable",
                    "r_squared": float(r_value**2),
                    "start_val": float(seg_data.iloc[0]),
                    "end_val": float(seg_data.iloc[-1])
                })
            last_bk = bk
            
        # 5. Abrupt Shock Detection (Anomaly detection on Remainder)
        # We use a robust threshold based on Median Absolute Deviation (MAD)
        mad = np.median(np.abs(remainder - np.median(remainder)))
        # Threshold: 3.5 * MAD, but at least 0.05 to avoid reporting noise/insignificant changes
        threshold = max(3.5 * mad, 0.05) if mad > 0 else 0.05
        
        shocks = []
        if threshold > 0:
            shock_indices = np.where(np.abs(remainder) > threshold)[0]
            for idx in shock_indices:
                shocks.append({
                    "date": trend.index[idx].strftime("%Y-%m-%d"),
                    "magnitude": float(remainder.iloc[idx]),
                    "severity": "Alta" if abs(remainder.iloc[idx]) > 2 * threshold else "Moderada"
                })

        # 6. Summary Results
        return {
            "trend": trend.tolist(),
            "seasonal": seasonal.tolist(),
            "remainder": remainder.tolist(),
            "dates": trend.index.strftime("%Y-%m-%d").tolist(),
            "breakpoints": [trend.index[bk-1].strftime("%Y-%m-%d") for bk in breakpoints_indices if 0 < bk < len(trend)],
            "segments": segments,
            "shocks": shocks,
            "magnitude": float(trend.iloc[-1] - trend.iloc[0]),
            "is_stable": abs(float(trend.iloc[-1] - trend.iloc[0])) < 0.05
        }
    except Exception as e:
        logger.error(f"BFAST Analysis Failed: {e}")
        return None

# ==========================================
# MODULE: REPORT GENERATOR
# ==========================================

def get_index_status_message(mode: str, stats: Dict, bfast: Dict) -> str:
    """
    Set de respuestas automatizadas: Generate diagnostic messages based on 
    statistical results (Trend, Significance, Moran's I, BFAST).
    """
    trend = stats.get('trend', 0)
    p_val = stats.get('p_value', 1.0)
    moran = stats.get('moran_i', 0)
    shocks = bfast.get('shocks', []) if bfast else []
    mag = bfast.get('magnitude', 0) if bfast else (trend / 100.0 if trend is not None else 0)
    
    is_sig = p_val < 0.05
    is_clustered = moran > 0.3
    
    # 1. Base Message by Mode
    base_msg = ""
    if mode == "Hydrology":
        if mag > 0.05: base_msg = "Expansión hídrica detectada." if is_sig else "Ligero aumento de humedad (no significativo)."
        elif mag < -0.05: base_msg = "Desecación crítica observada." if is_sig else "Tendencia a la reducción de agua (suave)."
        else: base_msg = "Niveles de agua estables."
    elif mode == "Vegetation":
        if mag > 0.05: base_msg = "Incremento vigoroso de biomasa." if is_sig else "Aumento leve de verdor."
        elif mag < -0.05: base_msg = "Pérdida severa de follaje/estrés." if is_sig else "Ligera disminución de biomasa."
        else: base_msg = "Vigor vegetal estable."
    elif mode == "Salinity":
        if mag > 0.05: base_msg = "Aumento preocupante de sales superficiales." if is_sig else "Tendencia al aumento de salinidad."
        elif mag < -0.05: base_msg = "Lavado de sales o mejora de sustrato." if is_sig else "Reducción moderada de salinidad."
        else: base_msg = "Niveles de salinidad estables."
    else:
        # Fallback for other indices
        if abs(mag) > 0.1: base_msg = f"Cambio estructural detectado en {get_index_name(mode)}."
        else: base_msg = f"Variaciones dentro del rango normal para {get_index_name(mode)}."

    # 2. Append Statistical Nuances
    reliability = "Alta fiabilidad estadística" if is_sig else "Baja fiabilidad estadística (ruido probable)"
    spatial = "Cambio espacialmente uniforme o aleatorio"
    if is_clustered: spatial = "Cambio focalizado en áreas específicas (Clustering detectado)"
    elif moran < -0.1: spatial = "Cambio disperso y fragmentado"

    full_status = f"{base_msg} {reliability}. {spatial}."
    
    if shocks:
        full_status += f" Se han identificado {len(shocks)} anomalías abruptas (choques) que podrían indicar eventos meteorológicos extremos."
        
    return full_status

def synthesize_wetland_health(analysis_results: Dict) -> Dict:
    """Detailed cross-analysis of all indices to determine the overall ecosystem health and recommendations."""
    # Extract trends & stats
    h = analysis_results.get('Hydrology', {}).get('stats', {})
    v = analysis_results.get('Vegetation', {}).get('stats', {})
    wq = analysis_results.get('WaterQuality', {}).get('stats', {})
    s = analysis_results.get('Salinity', {}).get('stats', {})
    sv = analysis_results.get('SoilVegetation', {}).get('stats', {})
    al = analysis_results.get('AlgaeBloom', {}).get('stats', {})
    wr = analysis_results.get('WaterRatio', {}).get('stats', {})

    h_t = h.get('trend', 0) or 0
    v_t = v.get('trend', 0) or 0
    wq_t = wq.get('trend', 0) or 0
    s_t = s.get('trend', 0) or 0
    sv_t = sv.get('trend', 0) or 0
    al_t = al.get('trend', 0) or 0
    wr_t = wr.get('trend', 0) or 0
    
    # Defaults
    conclusion = "Estabilidad Ecosistémica"
    severity = "Baja"
    details = "El sistema muestra variaciones dentro de los rangos normales de fluctuación estacional. "
    recs = [
        "Mantener el protocolo de monitoreo satelital estándar.",
        "Realizar una inspección visual de control en el próximo trimestre."
    ]

    # 1. Critical Water Loss & Vegetation Stress
    if h_t < -15 and v_t < -15:
        conclusion = "Deterioro Crítico por Desecación"
        severity = "Alta"
        details = f"Se observa una reducción simultánea de la lámina de agua ({h_t:.1f}%) y del vigor vegetal ({v_t:.1f}%). "
        if s_t > 15:
            details += "Este cuadro se ve agravado por un incremento en la salinidad superficial, indicando procesos de evaporación intensa."
        recs = [
            "Activar protocolo de emergencia por estrés hídrico.",
            "Realizar medición de conductividad eléctrica y pH en puntos críticos.",
            "Evaluar el estado de los canales de alimentación o fuentes de agua superficial."
        ]
    
    # 2. Eutrophication Risk
    elif wq_t > 20 and al_t > 15:
        conclusion = "Riesgo de Eutrofización / Bloom Algal"
        severity = "Alta"
        details = f"Incremento significativo de clorofila-a ({wq_t:.1f}%) y algas flotantes ({al_t:.1f}%) con baja renovación hídrica. "
        recs = [
            "Tomar muestras de agua para análisis de nitrógeno y fósforo.",
            "Identificar posibles fuentes de vertido de nutrientes en la cuenca aportante.",
            "Monitorear niveles de oxígeno disuelto para prevenir anoxia."
        ]

    # 3. Salinization dominance
    elif s_t > 20 and h_t < 0:
        conclusion = "Estrés por Salinización Superficial"
        severity = "Media"
        details = f"Detección de un aumento significativo en la firma de sales ({s_t:.1f}%) coincidente con una disminución de humedad. "
        recs = [
            "Validar en terreno la presencia de costras salinas.",
            "Analizar el impacto en la vegetación halófita local.",
            "Verificar intrusiones de aguas salobres si aplica a la zona."
        ]

    # 4. Recovery / Expansion
    elif h_t > 15 and v_t > 15:
        conclusion = "Fase de Recuperación e Inundación"
        severity = "Informativa"
        details = "Aumento de la disponibilidad hídrica propiciando una respuesta biológica positiva y expansión del área inundada."
        recs = [
            "Documentar la extensión máxima de la lámina de agua.",
            "Monitorear la colonización de nuevas áreas por vegetación hidrófila."
        ]

    # 5. Vegetation Stress without water loss
    elif v_t < -20 and h_t >= -5:
        conclusion = "Estrés Vegetativo (Posible Plaga o Quema)"
        severity = "Media"
        details = "Pérdida de vigor fotosintético no asociada directamente a la falta de agua superficial."
        recs = [
            "Investigar indicios de incendios o quemas de pastizales.",
            "Revisar presencia de especies invasoras o plagas forestales."
        ]

    # 6. Specific Water Quality Issue
    elif wq_t > 30 and al_t <= 5:
        conclusion = "Alteración de la Turbidez / Calidad de Agua"
        severity = "Media"
        details = "Cambios significativos en la coloración o turbidez del agua sin señales de florecimiento algal evidente."
        recs = [
            "Evaluar procesos de sedimentación por erosión aguas arriba.",
            "Revisar entradas de sedimentos tras eventos de lluvia intensos."
        ]

    # Add Moran's I context if consistent
    h_moran = h.get('moran_i', 0)
    if h_moran and h_moran > 0.6:
        details += " El patrón de cambio muestra una alta consistencia espacial, sugiriendo una tendencia estructural y no aleatoria."

    return {
        "conclusion": conclusion, 
        "details": details, 
        "severity": severity,
        "recommendations": recs,
        "trends": {"H": h_t, "V": v_t, "S": s_t}
    }

def download_image(url: str) -> io.BytesIO:
    """Download image from URL to BytesIO."""
    if not url: return None
    try:
        response = requests.get(url, timeout=60)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except Exception as e:
        logger.error(f"Error downloading image: {e}")
        return None

def get_vis_params(mode: str) -> Dict[str, Any]:
    """Get visualization parameters."""
    if mode == "Hydrology": return {'min': -1, 'max': 1, 'palette': ['FF0000', 'FFFFFF', '0000FF']}
    elif mode == "Vegetation": return {'min': 0, 'max': 0.8, 'palette': ['FF0000', 'FFFF00', '00FF00', '006400']}
    elif mode == "WaterQuality": return {'min': -0.1,  'max': 0.5, 'palette': ['0000FF', '00FFFF', 'FFFF00', 'FF0000']}
    elif mode == "SoilVegetation": return {'min': 0, 'max': 1, 'palette': ['FFFFFF', 'CE7E45', 'DF923D', 'F1B555', 'FCD163', '99B718', '74A901', '66A000', '529400', '3E8601', '207401', '056201', '004C00', '023B01', '012E01', '011D01', '011301']}
    elif mode == "AlgaeBloom": return {'min': -0.05, 'max': 0.2, 'palette': ['0000FF', '00FFFF', '00FF00', 'FFFF00', 'FF0000', '8B0000']}
    elif mode == "WaterRatio": return {'min': -1, 'max': 1, 'palette': ['FF0000', 'FFA500', 'FFFF00', 'FFFFFF', '00FFFF', '0000FF']}
    elif mode == "Salinity": return {'min': -1, 'max': 1, 'palette': ['0000FF', 'FFFFFF', '8B4513']} # Salty: Brown/Dry
    return {'min': 0, 'max': 1, 'palette': ['000000', 'FFFFFF']}

def get_index_name(mode: str) -> str:
    """Get the index name for a mode."""
    names = {'Hydrology': 'MNDWI', 'Vegetation': 'NDRE', 'WaterQuality': 'NDCI', 'SoilVegetation': 'SAVI', 'AlgaeBloom': 'FAI', 'WaterRatio': 'WRI', 'Salinity': 'NDSI'}
    return names.get(mode, mode)

def get_mode_description(mode: str) -> str:
    """Get description for each analysis mode."""
    descs = {
        'Hydrology': 'El índice MNDWI (Modified Normalized Difference Water Index) se utiliza para realzar cuerpos de agua abiertos y áreas de alta humedad. Valores positivos indican presencia de agua superficial, mientras que valores negativos representan suelo o vegetación seca.',
        'Vegetation': 'El índice NDRE (Normalized Difference Red Edge) es sensible al contenido de clorofila en la vegetación densa. Es fundamental para monitorear el vigor fotosintético y detectar estrés hídrico temprano en vegetación de humedal.',
        'WaterQuality': 'El índice NDCI (Normalized Difference Chlorophyll Index) permite estimar la concentración de clorofila-a en cuerpos de agua. Es un indicador clave del estado trófico y la posible presencia de fitoplancton en aguas lénticas.',
        'SoilVegetation': 'El índice SAVI (Soil Adjusted Vegetation Index) minimiza la influencia del brillo del suelo en el análisis de vegetación. Es ideal para humedales con cobertura vegetal dispersa o estacional.',
        'AlgaeBloom': 'El índice FAI (Floating Algae Index) detecta vegetación flotante y floraciones algales en la superficie del agua. Es crucial para identificar procesos de eutrofización y blooms de cianobacterias.',
        'WaterRatio': 'El índice WRI (Water Ratio Index) es un clasificador robusto para la discriminación entre superficies de agua y tierra. Valores > 1 indican una alta probabilidad de superficie acuática pura.',
        'Salinity': 'El índice NDSI (Normalized Difference Salinity Index) detecta la presencia de sales en la superficie del suelo. Es un indicador vital en zonas áridas o humedales con intrusión salina.'
    }
    return descs.get(mode, mode)

def create_legend_image(mode: str) -> io.BytesIO:
    """Create a legend image for the specific mode."""
    params = get_vis_params(mode)
    palette_hex = [f"#{c}" for c in params['palette']]
    fig, ax = plt.subplots(figsize=(6, 1))
    fig.subplots_adjust(bottom=0.5)
    cmap = mcolors.LinearSegmentedColormap.from_list("custom_cmap", palette_hex)
    norm = mcolors.Normalize(vmin=params['min'], vmax=params['max'])
    cb = fig.colorbar(plt.cm.ScalarMappable(norm=norm, cmap=cmap), cax=ax, orientation='horizontal')
    cb.set_label(f'Valor {get_index_name(mode)}')
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buffer.seek(0)
    return img_buffer

def create_temporal_chart(time_series: List[Dict], mode: str) -> io.BytesIO:
    """Create a temporal chart for a specific mode."""
    dates = [point['date'] for point in time_series if point.get('value') is not None]
    values = [point['value'] for point in time_series if point.get('value') is not None]
    if not dates: return None
    
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.plot(dates, values, marker='o', linewidth=2, markersize=4)
    ax.set_title(f'Serie Temporal - {mode}', fontsize=14, fontweight='bold')
    ax.set_xlabel('Fecha', fontsize=10)
    ax.set_ylabel('Valor del Índice', fontsize=10)
    ax.grid(True, alpha=0.3)
    plt.xticks(rotation=45, ha='right')
    
    if len(dates) > 20:
        nth = len(dates) // 10
        for i, label in enumerate(ax.xaxis.get_ticklabels()):
            if i % nth != 0: label.set_visible(False)
            
    plt.tight_layout()
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buffer.seek(0)
    return img_buffer

def generate_wetland_report(wetland_name: str, wetland_metadata: Dict, analysis_results: Dict, start_date: str, end_date: str) -> io.BytesIO:
    """Generate a high-end, professional technical sheet (Ficha Técnica) for wetland monitoring."""
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # --- HEADER SECTION (3 levels) ---
    header_table = doc.add_table(rows=3, cols=1)
    header_table.style = 'Table Grid'
    header_table.autofit = True
    
    c0 = header_table.rows[0].cells[0]
    c0.paragraphs[0].text = "CUADRO DE FICHAS TECNICAS OFICIAL"
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c0.paragraphs[0].runs[0].font.bold = True
    c0.paragraphs[0].runs[0].font.size = Pt(12)
    c0.paragraphs[0].style.font.color.rgb = RGBColor(0, 0, 0)
    # Set background color (light grey/greenish as per reference)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9EAD3') # Light green
    c0._tc.get_or_add_tcPr().append(shading_elm)

    c1 = header_table.rows[1].cells[0]
    p1 = c1.paragraphs[0]
    p1.add_run("PROYECTO: ").bold = True
    p1.add_run(f"MONITOREO SATELITAL AVANZADO Y DIAGNÓSTICO ECOLÓGICO: {wetland_name.upper()}")
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.runs[0].font.size = Pt(9)

    c2 = header_table.rows[2].cells[0]
    p2 = c2.paragraphs[0]
    p2.add_run("ENTIDAD: ").bold = True
    p2.add_run(wetland_metadata.get('entity', 'SISTEMA DE MONITOREO AMBIENTAL GEOINT'))
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # --- METADATA GRID (3 Columns) ---
    doc.add_heading('DESCRIPCION DE FICHA TECNICA DE MONITOREO', level=3).alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_grid = doc.add_table(rows=8, cols=3)
    meta_grid.style = 'Table Grid'
    
    def set_cell_text(row, col, label, val):
        p = meta_grid.rows[row].cells[col].paragraphs[0]
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(val))
        p.runs[0].font.size = Pt(8)
        if len(p.runs) > 1: p.runs[1].font.size = Pt(8)

    # Labels Row 0, 2, 4, 6 | Values Row 1, 3, 5, 7
    set_cell_text(0, 0, "PAÍS", "")
    set_cell_text(1, 0, "", wetland_metadata.get('department', 'CHILE'))
    set_cell_text(0, 1, "CARACTERISTICAS", "")
    set_cell_text(1, 1, "", "HUMEDAL PRIORITARIO")
    set_cell_text(0, 2, "DESIGNACIÓN", "")
    set_cell_text(1, 2, "", wetland_metadata.get('code', 'W-01'))

    set_cell_text(2, 0, "PROVINCIA", "")
    set_cell_text(3, 0, "", wetland_metadata.get('province', 'AREA DE ESTUDIO'))
    set_cell_text(2, 1, "ESTABLECIDA POR", "")
    set_cell_text(3, 1, "", "WETLAND MONITOR AI")
    set_cell_text(2, 2, "PERIODO", "")
    set_cell_text(3, 2, "", f"{start_date} / {end_date}")

    set_cell_text(4, 0, "DISTRITO", "")
    set_cell_text(5, 0, "", wetland_name)
    set_cell_text(4, 1, "COORDENADAS CENTROIDE", "")
    set_cell_text(5, 1, "", wetland_metadata.get('coordinates', 'N/A'))
    set_cell_text(4, 2, "DATUM", "")
    set_cell_text(5, 2, "", "WGS-84")

    set_cell_text(6, 0, "UBICACIÓN", "")
    set_cell_text(7, 0, "", "COORDENADAS GEOGRÁFICAS")
    set_cell_text(6, 1, "ESTADO GLOBAL", "")
    set_cell_text(7, 1, "", synthesize_wetland_health(analysis_results)['conclusion'])
    set_cell_text(6, 2, "PRECISIÓN", "")
    set_cell_text(7, 2, "", "SENTINEL-2 (10M)")

    doc.add_paragraph()

    # --- DETAILED ANALYSIS PER INDEX ---
    doc.add_heading('DESCRIPCION Y ANALISIS DE INDICADORES SATELLITALES', level=3).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    modes = ['Hydrology', 'Vegetation', 'WaterQuality', 'SoilVegetation', 'AlgaeBloom', 'WaterRatio', 'Salinity']
    for mode in modes:
        if mode not in analysis_results: continue
        res = analysis_results[mode]
        stats = res.get('stats', {})
        
        # --- FILTER: Skip indices with no data or zero value ---
        curr = stats.get('current')
        if curr is None or curr == 0:
            logger.info(f"Skipping index {mode} in report: No signal or data (value: {curr})")
            continue
            
        bfast = res.get('bfast', {})
        
        # --- ENCAPSULATING BOX (RECUADRO) FOR INDEX ---
        # Create a single-cell table to act as a container with borders
        container_table = doc.add_table(rows=1, cols=1)
        container_table.style = 'Table Grid'
        container_cell = container_table.rows[0].cells[0]
        
        # Everything from here on goes into container_cell
        container_cell.paragraphs[0].text = f'ÍNDICE: {get_index_name(mode)}'
        container_cell.paragraphs[0].runs[0].font.bold = True
        container_cell.paragraphs[0].runs[0].font.size = Pt(11)
        container_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Background for header part of the box
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F3F3F3')
        container_cell._tc.get_or_add_tcPr().append(shd)

        # 1. Description Párrafo
        p_desc = container_cell.add_paragraph()
        p_desc.add_run("DESCRIPCION: ").bold = True
        p_desc.add_run(get_mode_description(mode))
        p_desc.runs[0].font.size = Pt(8)
        p_desc.runs[1].font.size = Pt(8)

        # 2. Diagnosis Párrafo
        status_msg = get_index_status_message(mode, stats, bfast)
        p_diag = container_cell.add_paragraph()
        p_diag.add_run("DIAGNOSTICO ESPECIFICO: ").bold = True
        p_diag.add_run(status_msg)
        p_diag.runs[0].font.size = Pt(8)
        p_diag.runs[1].font.size = Pt(8)

        # 3. Stats Row
        curr = stats.get('current')
        prev = stats.get('last')
        trend = stats.get('trend')
        
        curr_str = f"{curr:.4f}" if curr is not None else "SIN DATOS"
        prev_str = f"{prev:.4f}" if prev is not None else "SIN DATOS"
        trend_str = f"{trend:+.1f}%" if trend is not None else "N/A"
        
        p_stats = container_cell.add_paragraph()
        p_stats.add_run("VALOR ACTUAL: ").bold = True
        p_stats.add_run(f"{curr_str}  ")
        p_stats.add_run("VALOR INICIAL: ").bold = True
        p_stats.add_run(f"{prev_str}  ")
        p_stats.add_run("VARIACIÓN: ").bold = True
        run_v = p_stats.add_run(trend_str)
        if trend and trend > 10: run_v.font.color.rgb = RGBColor(34, 197, 94)
        if trend and trend < -10: run_v.font.color.rgb = RGBColor(239, 68, 68)
        for r in p_stats.runs: r.font.size = Pt(8)

        # 4. Maps (if available) - Inner Table for side-by-side
        maps = res.get('maps', {})
        img_start = download_image(maps.get('start_year', {}).get('thumb_url'))
        img_end = download_image(maps.get('end_year', {}).get('thumb_url'))

        if img_start or img_end:
            p_map_label = container_cell.add_paragraph()
            p_map_label.add_run("COMPARATIVA ESPACIAL:").bold = True
            p_map_label.runs[0].font.size = Pt(8)
            
            p_imgs = container_cell.add_paragraph()
            p_imgs.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_start:
                p_imgs.add_run().add_picture(img_start, width=Inches(2.5))
                p_imgs.add_run("   ")
            if img_end:
                p_imgs.add_run().add_picture(img_end, width=Inches(2.5))

        # 5. Temporal Chart
        ts = res.get('time_series', [])
        if ts:
            chart = create_temporal_chart(ts, mode)
            if chart:
                p_chart_label = container_cell.add_paragraph()
                p_chart_label.add_run("COMPORTAMIENTO TEMPORAL:").bold = True
                p_chart_label.runs[0].font.size = Pt(8)
                
                p_chart_img = container_cell.add_paragraph()
                p_chart_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_chart_img.add_run().add_picture(chart, width=Inches(5.5))

        # Each index gets exactly one page
        doc.add_page_break()

    # --- FINAL AUTOMATED CONCLUSIONS PAGE ---
    doc.add_heading('CONCLUSIONES TÉCNICAS DEL MONITOREO', level=3).alignment = WD_ALIGN_PARAGRAPH.CENTER
    health = synthesize_wetland_health(analysis_results)
    
    conc_table = doc.add_table(rows=1, cols=1)
    conc_table.style = 'Table Grid'
    conc_cell = conc_table.rows[0].cells[0]
    
    # Background for conclusion header
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'E2EFE0')
    conc_cell._tc.get_or_add_tcPr().append(shd)
    
    p_conc = conc_cell.paragraphs[0]
    p_conc.add_run("DIAGNÓSTICO SINTÉTICO: ").bold = True
    p_conc.add_run(health['conclusion'].upper())
    p_conc.runs[1].font.color.rgb = RGBColor(255, 0, 0) if health['severity'] == "Alta" else RGBColor(0, 0, 0)
    
    p_det = conc_cell.add_paragraph()
    p_det.add_run("SÍNTESIS NARRATIVA: ").bold = True
    p_det.add_run(health['details'])
    p_det.runs[0].font.size = Pt(10)
    p_det.runs[1].font.size = Pt(10)
    
    p_sev = conc_cell.add_paragraph()
    p_sev.add_run("NIVEL DE CRITICIDAD: ").bold = True
    p_sev.add_run(health['severity'])
    p_sev.runs[1].font.bold = True
    
    doc.add_paragraph()
    p_rec = doc.add_paragraph()
    p_rec.add_run("RECOMENDACIONES:").bold = True
    for i, rec in enumerate(health.get('recommendations', []), 1):
        p_rec.add_run(f"\n{i}. {rec}")
    
    # Save
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# ==========================================
# APP & GEE CONFIGURATION
# ==========================================

app = FastAPI(title="GEOINT Wetland Monitor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# GEE Initialization (Attempt global init, but handle failure gracefully)
try:
    # Try global init without project first (for backward compatibility/local auth)
    ee.Initialize()
    logger.info("Google Earth Engine Initialized globally")
except Exception:
    logger.warning("Global GEE init failed. Will attempt per-request initialization.")

def ensure_ee_initialized(project_id: str = None):
    """Ensure EE is initialized with a project ID."""
    try:
        # Check if already initialized
        ee.Projection('EPSG:4326') 
    except Exception:
        # Not initialized or needs project
        try:
            p = project_id or os.getenv("GEE_PROJECT_ID", "ee-wetlandmonitor")
            ee.Initialize(project=p)
            logger.info(f"GEE Initialized with project: {p}")
        except Exception as e:
            logger.error(f"Failed to initialize GEE: {e}")
            raise ValidationError(f"Earth Engine initialization failed: {e}")

class AnalysisRequest(BaseModel):
    geojson: Dict[str, Any]
    startDate: str
    endDate: str
    projectId: str = None
    mode: str = "Hydrology"

def normalize_index_value(value, mode):
    """Normalize index values for consistent charting."""
    if value is None: return None
    
    ranges = {
        'Hydrology': (-1, 1),      # MNDWI
        'Vegetation': (0, 0.8),    # NDRE
        'WaterQuality': (-0.1, 0.5), # NDCI
        'SoilVegetation': (0, 1),  # SAVI
        'AlgaeBloom': (-0.05, 0.2), # FAI
        'WaterRatio': (-1, 1),      # WRI
        'Salinity': (-1, 1)         # NDSI
    }
    
    if mode == 'WaterRatio':
        # Logarithmic normalization for WRI
        if value <= 0: return -1
        try:
            log_val = math.log10(value)
            return max(-1, min(1, log_val))
        except: return -1
        
    min_v, max_v = ranges.get(mode, (-1, 1))
    return max(min_v, min(max_v, value))

def get_sentinel_data(aoi, start_date, end_date, mode):
    """Get and process Sentinel-2 data with scaling and cloud masking."""
    
    def mask_clouds(img):
        cloud_bit_mask = 1 << 10
        cirrus_bit_mask = 1 << 11
        qa = img.select('QA60')
        # SCL (Scene Classification Layer) for more robust masking if available
        # 3: Cloud Shadows, 8: Cloud Medium Prob, 9: Cloud High Prob, 10: Cirrus
        mask = qa.bitwiseAnd(cloud_bit_mask).eq(0).And(
               qa.bitwiseAnd(cirrus_bit_mask).eq(0))
        
        # Scale bands from 0-10000 to 0-1 (IMPORTANT for SAVI/FAI/NDCI)
        return img.updateMask(mask).divide(10000).copyProperties(img, ['system:time_start'])

    s2_col = (ee.ImageCollection("COPERNICUS/S2_SR_HARMONIZED")
              .filterBounds(aoi)
              .filterDate(start_date, end_date)
              .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 20))
              .map(mask_clouds))

    if mode == "Hydrology":
        # MNDWI (B3 Green, B11 SWIR)
        def add_mndwi(img):
            mndwi = img.normalizedDifference(['B3', 'B11']).rename('Value')
            return img.addBands(mndwi).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(add_mndwi)

    elif mode == "Vegetation":
        # NDRE (B8 NIR, B5 RedEdge)
        def add_ndre(img):
            ndre = img.normalizedDifference(['B8', 'B5']).rename('Value')
            return img.addBands(ndre).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(add_ndre)

    elif mode == "WaterQuality":
        # NDCI (B5 RedEdge, B4 Red)
        def calc_ndci(img):
            ndci = img.normalizedDifference(['B5', 'B4']).rename('Value')
            return img.addBands(ndci).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(calc_ndci)

    elif mode == "SoilVegetation":
        # SAVI ((B8 - B4) / (B8 + B4 + 0.5)) * 1.5
        def calc_savi(img):
            savi = img.expression(
                '((NIR - RED) / (NIR + RED + 0.5)) * 1.5',
                {'NIR': img.select('B8'), 'RED': img.select('B4')}
            ).rename('Value')
            return img.addBands(savi).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(calc_savi)
        
    elif mode == "AlgaeBloom":
        # FAI (B8 - (B4 + (B11-B4) * (832.8-664.6)/(1613.7-664.6)))
        def calc_fai(img):
            fai = img.expression(
                'NIR - (RED + (SWIR - RED) * 0.177)',
                {'NIR': img.select('B8'), 'RED': img.select('B4'), 'SWIR': img.select('B11')}
            ).rename('Value')
            return img.addBands(fai).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(calc_fai)
    
    elif mode == "WaterRatio":
        # WRI (Green + Red) / (NIR + SWIR) -> (B3 + B4) / (B8 + B11)
        def calc_wri(img):
            wri = img.expression(
                '(GREEN + RED) / (NIR + SWIR)',
                {'GREEN': img.select('B3'), 'RED': img.select('B4'), 'NIR': img.select('B8'), 'SWIR': img.select('B11')}
            ).rename('Value')
            return img.addBands(wri).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(calc_wri)
        
    elif mode == "Salinity":
        # NDSI (SWIR1 - NIR) / (SWIR1 + NIR) -> (B11 - B8) / (B11 + B8)
        def calc_ndsi(img):
            ndsi = img.normalizedDifference(['B11', 'B8']).rename('Value')
            return img.addBands(ndsi).select('Value').copyProperties(img, ['system:time_start'])
        return s2_col.map(calc_ndsi)

    return ee.ImageCollection([])

def analyze_period(aoi, start_date, end_date, mode):
    """Analyze a specific period for time series data."""
    col = get_sentinel_data(aoi, start_date, end_date, mode)
    
    def reduce_img(img):
        date = img.date().format("YYYY-MM-dd")
        mean_val = img.reduceRegion(
            reducer=ee.Reducer.median(), # Median is robust to outliers
            geometry=aoi,
            scale=10, # Reverted to native 10m resolution for maximum detail
            maxPixels=1e9
        ).get('Value')
        return ee.Feature(None, {'date': date, 'value': mean_val})
        
    features = col.map(reduce_img).filter(ee.Filter.notNull(['value'])).getInfo()['features']
    return [{'date': f['properties']['date'], 'value': f['properties']['value']} for f in features]

def generate_map_url(aoi, start, end, mode):
    """Generate map tile URLs for RGB and metric visualization."""
    col = get_sentinel_data(aoi, start, end, mode)
    latest = col.median().clip(aoi)
    
    # RGB visualization logic
    if mode == "Hydrology":
        s2_col = (ee.ImageCollection("COPERNICUS/S2_SR_HARMONIZED")
                  .filterBounds(aoi).filterDate(start, end)
                  .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 30)))
        s2_rgb = s2_col.median().clip(aoi).divide(10000)
        rgb_vis = {'min': 0, 'max': 0.3, 'bands': ['B4', 'B3', 'B2']}
        rgb_map = s2_rgb.getMapId(rgb_vis)
    else:
        rgb_vis = {'min': 0, 'max': 0.3, 'bands': ['B4', 'B3', 'B2']}
        # Fallback to S2 collection for RGB if 'latest' only has Value band
        s2_col_rgb = (ee.ImageCollection("COPERNICUS/S2_SR_HARMONIZED")
                  .filterBounds(aoi).filterDate(start, end)
                  .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 30)))
        rgb_map = s2_col_rgb.median().clip(aoi).divide(10000).getMapId(rgb_vis)
    
    metric_vis = get_vis_params(mode)
    metric_map = latest.select('Value').getMapId(metric_vis)
    
    # Generate thumbnail URL for report
    thumb_params = metric_vis.copy()
    thumb_params['dimensions'] = 350
    thumb_params['region'] = aoi
    thumb_params['format'] = 'png'
    try:
        thumb_url = latest.select('Value').getThumbURL(thumb_params)
    except Exception:
        thumb_url = None

    return {
        "rgb": rgb_map['tile_fetcher'].url_format,
        "metric": metric_map['tile_fetcher'].url_format,
        "thumb_url": thumb_url
    }

def perform_single_analysis(request, mode):
    """Perform robust analysis for a single mode."""
    try:
        ensure_ee_initialized(request.projectId)
        logger.info(log_process_stage('', mode, 'processing'))
        start_obj, end_obj = validate_date_range(request.startDate, request.endDate)
        aoi = validate_geometry(request.geojson)
        
        # Auto-expand temporal window if insufficient data (Option 1)
        current_data = analyze_period(aoi, request.startDate, request.endDate, mode)
        expansions = 0
        current_start_obj = start_obj

        while (not current_data or len([d for d in current_data if d.get('value') is not None]) < 1) and expansions < 4:
            expansions += 1
            current_start_obj = current_start_obj - relativedelta(months=3)
            logger.info(f"[{mode}] Insufficient data. Expanding temporal window backwards to {current_start_obj.strftime('%Y-%m-%d')}")
            current_data = analyze_period(aoi, current_start_obj.strftime("%Y-%m-%d"), request.endDate, mode)

        if not current_data or len([d for d in current_data if d.get('value') is not None]) < 1:
            logger.warning(f"Still insufficient data for {mode} after expanding 1 year backwards. Using whatever is available.")
            
        coverage = validate_temporal_coverage(current_data)
        if not coverage['valid']:
            raise ValidationError(f"Temporal coverage issue: {coverage['reason']}")
            
        current_stats = calculate_robust_statistics(current_data)
        current_data_flagged = detect_outliers(current_data)
        outlier_count = sum(1 for d in current_data_flagged if d.get('is_outlier'))
        
        # Baseline logic: If range > 2 years, compare against the beginning of the period (Historical Baseline)
        total_range_days = (end_obj - start_obj).days
        if total_range_days > 730:
            last_start = request.startDate
            last_end = (start_obj + relativedelta(years=1)).strftime("%Y-%m-%d")
            logger.info(f"[{mode}] Using Historical Baseline for statistics (Long range: {total_range_days} days)")
            
            last_data = analyze_period(aoi, last_start, last_end, mode)
            # Auto-expansion Forward for Historical Baseline if data is missing (e.g. 2016 clouds)
            b_exp = 0
            while (not last_data or len([d for d in last_data if d.get('value') is not None]) < 1) and b_exp < 4:
                b_exp += 1
                last_end_obj = datetime.strptime(last_end, "%Y-%m-%d") + relativedelta(months=3)
                last_end = last_end_obj.strftime("%Y-%m-%d")
                logger.info(f"[{mode}] Insufficient baseline data. Expanding baseline window forward to {last_end}")
                last_data = analyze_period(aoi, last_start, last_end, mode)
        else:
            last_start = (current_start_obj - relativedelta(years=1)).strftime("%Y-%m-%d")
            last_end = (end_obj - relativedelta(years=1)).strftime("%Y-%m-%d")
            last_data = analyze_period(aoi, last_start, last_end, mode)
            
            # Expansion for Standard Baseline
            b_exp = 0
            while (not last_data or len([d for d in last_data if d.get('value') is not None]) < 1) and b_exp < 4:
                b_exp += 1
                last_start_obj = datetime.strptime(last_start, "%Y-%m-%d") - relativedelta(months=3)
                last_start = last_start_obj.strftime("%Y-%m-%d")
                logger.info(f"[{mode}] Insufficient standard baseline data. Expanding backwards to {last_start}")
                last_data = analyze_period(aoi, last_start, last_end, mode)

        trend_stats = calculate_trend_statistics(current_data_flagged, last_data) or {}
        
        # Start/End Year Maps
        start_year_end = (start_obj + relativedelta(years=1)).strftime("%Y-%m-%d")
        maps_start = generate_map_url(aoi, request.startDate, start_year_end, mode)
        end_year_start = (end_obj - relativedelta(years=1)).strftime("%Y-%m-%d")
        maps_end = generate_map_url(aoi, end_year_start, request.endDate, mode)
        
        maps = {
            "rgb": maps_end["rgb"],
            "metric": maps_end["metric"],
            "start_year": maps_start,
            "end_year": maps_end
        }
        
        result = {
            "mode": mode,
            "stats": {
                "current": current_stats['median'],
                "current_mean": current_stats['mean'],
                "current_std": current_stats['std'],
                "last": trend_stats.get('previous_median', 0),
                "trend": trend_stats.get('trend_percent', 0),
                "p_value": trend_stats.get('p_value', 1.0),
                "is_significant": trend_stats.get('is_significant', False),
                "moran_i": calculate_spatial_consistency(get_sentinel_data(aoi, end_year_start, request.endDate, mode).median().clip(aoi), aoi),
                "outlier_count": outlier_count,
                "data_count": current_stats['count'],
                "cv": current_stats['cv'],
                "baseline_period": f"{last_start} a {last_end}"
            },
            "time_series": current_data_flagged,
            "maps": maps,
            "coverage": coverage,
            "bfast": apply_bfast_analysis(current_data_flagged, mode)
        }
        
        # Normalize time series for display
        normalized_series = []
        for point in current_data_flagged:
            p = point.copy()
            p['value_raw'] = point['value']
            p['value'] = normalize_index_value(point['value'], mode)
            normalized_series.append(p)
        result['time_series'] = normalized_series
        
        # Internal Auditing for Validation
        try:
            audit_entry = {
                "timestamp": datetime.now().isoformat(),
                "mode": mode,
                "project": request.projectId,
                "median_value": current_stats['median'],
                "data_points": current_stats['count'],
                "sample_points": current_data_flagged[:3]
            }
            with open("audit_indices.json", "a") as af:
                af.write(json.dumps(audit_entry) + "\n")
        except Exception as ae:
            logger.error(f"Audit log failed: {ae}")

        logger.info(log_process_stage('', mode, 'completed'))
        return result
        
    except Exception as e:
        logger.error(f"{mode} error: {e}")
        return create_error_response(e, mode)

@app.get("/debug/indices")
async def debug_indices():
    """Internal diagnostic endpoint to verify real GEE calculations."""
    try:
        # Internal test configuration
        ensure_ee_initialized(None)
        
        aoi = ee.Geometry.Point([-73.0928, -41.6793]).buffer(500).bounds()
        start = "2024-01-01"
        end = "2024-03-12"
        class MockRequest:
            def __init__(self, geojson, startDate, endDate, projectId):
                self.geojson = geojson
                self.startDate = startDate
                self.endDate = endDate
                self.projectId = projectId
                self.mode = "Hydrology"

        request = MockRequest(aoi.getInfo(), start, end, project_id)
        result = perform_single_analysis(request, "Hydrology")
        
        return {
            "status": "success",
            "project": project_id,
            "mode": mode,
            "analysis_result": result
        }
    except Exception as e:
        logger.error(f"Debug endpoint failed: {e}")
        return {"status": "error", "message": str(e)}

@app.post("/analyze")
async def analyze(request: AnalysisRequest, authorization: str = Header(None)):
    if not authorization: raise HTTPException(401, "Missing Token")
    token = authorization.split(" ")[1]
    
    try:
        from google.oauth2.credentials import Credentials
        creds = Credentials(token)
        if request.projectId: ee.Initialize(creds, project=request.projectId)
        else: ee.Initialize(creds)
        
        res = perform_single_analysis(request, request.mode)
        if not res or 'error' in res: 
            raise HTTPException(500, str(res.get('error', 'Unknown Error')))
            
        formatted_series = []
        for d in res['time_series']:
            formatted_series.append({
                "date": d['date'],
                "value": d['value'],
                "metric_name": get_index_name(request.mode)
            })
        formatted_series.sort(key=lambda x: x['date'])
        
        return {
            "status": "success",
            "data": {
                "time_series": formatted_series,
                "summary": {
                    "current_avg": res['stats']['current'],
                    "last_year_avg": res['stats']['last'],
                    "trend": res['stats']['trend'],
                    "mode": request.mode
                },
                "maps": res['maps']
            }
        }
    except Exception as e:
        raise HTTPException(500, detail=str(e))

@app.post("/analyze-all")
async def analyze_all(request: AnalysisRequest, authorization: str = Header(None)):
    if not authorization: raise HTTPException(401, "Missing Token")
    
    try:
        token = authorization.split(" ")[1]
        from google.oauth2.credentials import Credentials
        creds = Credentials(token)
        
        if request.projectId: 
            ee.Initialize(creds, project=request.projectId)
        else: 
            ee.Initialize(creds)
        
        results = {}
        modes = ["Hydrology", "Vegetation", "WaterQuality", "SoilVegetation", "AlgaeBloom", "WaterRatio", "Salinity"]
        
        # Parallel Execution of Analysis Modes
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(modes)) as executor:
            # Map modes to futures
            future_to_mode = {executor.submit(perform_single_analysis, request, m): m for m in modes}
            
            for future in concurrent.futures.as_completed(future_to_mode):
                m = future_to_mode[future]
                try:
                    results[m] = future.result()
                    logger.info(f"Parallel Task {m}: Success")
                except Exception as exc:
                    logger.error(f"Parallel Task {m} generated an exception: {exc}")
                    results[m] = create_error_response(exc, m)
            
        logger.info(log_process_stage('', None, 'final'))
        return {"status": "success", "data": results}
    except Exception as e:
        logger.error(f"Analyze-all Error: {e}")
        raise HTTPException(500, detail=f"Backend Error: {str(e)}")

@app.post("/process-spatial-file")
async def process_spatial_file(file: UploadFile = File(...)):
    """Process uploaded KML, KMZ, or SHP (zipped) files and return GeoJSON."""
    temp_dir = tempfile.mkdtemp()
    try:
        file_path = os.path.join(temp_dir, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Determine file type and process
        if file.filename.lower().endswith('.zip'):
            # Assume it's a Zipped Shapefile
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Find the .shp file
            shp_files = [f for f in os.listdir(temp_dir) if f.endswith('.shp')]
            if not shp_files:
                raise HTTPException(400, "No .shp file found in ZIP")
            gdf = gpd.read_file(os.path.join(temp_dir, shp_files[0]))
            
        elif file.filename.lower().endswith(('.kml', '.kmz')):
            # KMZ is a zipped KML. Unzip it first for robust processing.
            if file.filename.lower().endswith('.kmz'):
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                # Find the main KML file (usually doc.kml)
                kml_files = [f for f in os.listdir(temp_dir) if f.lower().endswith('.kml')]
                if not kml_files:
                    raise HTTPException(400, "No valid KML file found inside KMZ")
                file_path = os.path.join(temp_dir, kml_files[0])
                logger.info(f"Extracted KML from KMZ: {kml_files[0]}")

            # Robust KML reading (Multi-layer support)
            try:
                gdf = gpd.read_file(file_path)
                if gdf.empty:
                    layers = fiona.listlayers(file_path)
                    all_gdfs = []
                    for layer in layers:
                        try:
                            l_gdf = gpd.read_file(file_path, layer=layer)
                            if not l_gdf.empty: all_gdfs.append(l_gdf)
                        except: continue
                    if all_gdfs:
                        gdf = gpd.GeoDataFrame(pd.concat(all_gdfs, ignore_index=True))
            except Exception as e:
                logger.warning(f"Standard KML read failed, scanning layers: {e}")
                layers = fiona.listlayers(file_path)
                all_gdfs = []
                for layer in layers:
                    try:
                        l_gdf = gpd.read_file(file_path, layer=layer)
                        if not l_gdf.empty: all_gdfs.append(l_gdf)
                    except: continue
                if all_gdfs:
                    gdf = gpd.GeoDataFrame(pd.concat(all_gdfs, ignore_index=True))
                else:
                    raise HTTPException(400, f"No se pudo procesar el KML/KMZ: {str(e)}")
            
        else:
            raise HTTPException(400, "Unsupported file format. Use .kml, .kmz, or .zip (for SHP)")

        if gdf.empty:
            raise HTTPException(400, "The uploaded file contains no valid geometry")

        # Convert to WGS84
        if gdf.crs and gdf.crs.to_epsg() != 4326:
            gdf = gdf.to_crs(epsg=4326)
        
        # Merge all geometries into a single polygon/multipolygon
        combined_geom = gdf.geometry.unary_union
        
        # Force 2D coordinates (GEE can fail with 3D from KML)
        combined_geom = transform(to_2d, combined_geom)
        
        # Convert to GeoJSON
        feature = {
            "type": "Feature",
            "geometry": json.loads(gpd.GeoSeries([combined_geom]).to_json())['features'][0]['geometry'],
            "properties": {"name": file.filename}
        }
        
        # Calculate BBox
        bounds = combined_geom.bounds # (minx, miny, maxx, maxy)
        
        return {
            "status": "success",
            "geojson": feature,
            "bbox": [bounds[0], bounds[1], bounds[2], bounds[3]]
        }

    except Exception as e:
        logger.error(f"Error processing spatial file: {e}")
        raise HTTPException(500, detail=str(e))
    finally:
        # ignore_errors=True is important for Windows where files might be temporarily locked
        shutil.rmtree(temp_dir, ignore_errors=True)

@app.post("/generate-report")
async def generate_report(request: dict):
    try:
        wetland_name = request.get('wetland_name', 'Humedal Desconocido')
        logger.info(f"Generating report for: {wetland_name}")
        
        wetland_metadata = request.get('wetland_metadata', {})
        analysis_results = request.get('analysis_results', {})
        start_date = request.get('start_date', '')
        end_date = request.get('end_date', '')
        
        if not analysis_results:
            logger.warning("No analysis results provided for report generation")
            
        logger.debug(f"Analysis results keys: {list(analysis_results.keys()) if analysis_results else 'None'}")
        
        doc_buffer = generate_wetland_report(wetland_name, wetland_metadata, analysis_results, start_date, end_date)
        filename = f"Reporte_{wetland_name.replace(' ', '_')}_{end_date}.docx"
        
        logger.info(f"Report generated successfully: {filename}")
        
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Report Error: {e}")
        logger.error(f"Traceback: {error_trace}")
        raise HTTPException(500, detail=f"Report generation failed: {str(e)}")

@app.get("/")
def read_root():
    return {"status": "Backend running", "service": "Wetland Monitor AI"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
