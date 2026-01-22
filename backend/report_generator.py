"""
Word Report Generator for Wetland Analysis
Generates comprehensive .docx reports with analysis results, statistics, charts, and maps.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
from typing import Dict, Any, List
import numpy as np
import requests

def download_image(url: str) -> io.BytesIO:
    """Download image from URL to BytesIO."""
    if not url:
        return None
    try:
        response = requests.get(url, timeout=60)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except Exception as e:
        print(f"Error downloading image: {e}")
        return None

def get_vis_params(mode: str) -> Dict[str, Any]:
    """Get visualization parameters (mirrors main.py logic)."""
    if mode == "Hydrology": return {'min': -1, 'max': 1, 'palette': ['FF0000', 'FFFFFF', '0000FF']} # MNDWI
    elif mode == "Vegetation": return {'min': 0, 'max': 0.8, 'palette': ['FF0000', 'FFFF00', '00FF00', '006400']} # NDRE
    elif mode == "WaterQuality": return {'min': -0.1,  'max': 0.5, 'palette': ['0000FF', '00FFFF', 'FFFF00', 'FF0000']} # NDCI
    elif mode == "SoilVegetation": return {'min': 0, 'max': 1, 'palette': ['FFFFFF', 'CE7E45', 'DF923D', 'F1B555', 'FCD163', '99B718', '74A901', '66A000', '529400', '3E8601', '207401', '056201', '004C00', '023B01', '012E01', '011D01', '011301']} # SAVI
    elif mode == "AlgaeBloom": return {'min': -0.05, 'max': 0.2, 'palette': ['0000FF', '00FFFF', '00FF00', 'FFFF00', 'FF0000', '8B0000']} # FAI
    elif mode == "WaterRatio": return {'min': -1, 'max': 1, 'palette': ['FF0000', 'FFA500', 'FFFF00', 'FFFFFF', '00FFFF', '0000FF']} # WRI
    return {'min': 0, 'max': 1, 'palette': ['000000', 'FFFFFF']}

def create_legend_image(mode: str) -> io.BytesIO:
    """Create a legend image for the specific mode."""
    params = get_vis_params(mode)
    vis_min = params['min']
    vis_max = params['max']
    palette_hex = [f"#{c}" for c in params['palette']]
    
    fig, ax = plt.subplots(figsize=(6, 1))
    fig.subplots_adjust(bottom=0.5)
    
    cmap = mcolors.LinearSegmentedColormap.from_list("custom_cmap", palette_hex)
    norm = mcolors.Normalize(vmin=vis_min, vmax=vis_max)
    
    cb = fig.colorbar(plt.cm.ScalarMappable(norm=norm, cmap=cmap),
                      cax=ax, orientation='horizontal')
    cb.set_label(f'Valor {get_index_name(mode)}')
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buffer.seek(0)
    
    return img_buffer

def create_temporal_chart(time_series: List[Dict], mode: str) -> io.BytesIO:
    """
    Create a temporal chart for a specific mode.
    """
    # Extract dates and values
    dates = [point['date'] for point in time_series if point.get('value') is not None]
    values = [point['value'] for point in time_series if point.get('value') is not None]
    
    if not dates:
        return None
    
    # Create figure
    fig, ax = plt.subplots(figsize=(10, 4))
    
    # Plot data
    ax.plot(dates, values, marker='o', linewidth=2, markersize=4)
    
    # Styling
    ax.set_title(f'Serie Temporal - {mode}', fontsize=14, fontweight='bold')
    ax.set_xlabel('Fecha', fontsize=10)
    ax.set_ylabel('Valor del Índice', fontsize=10)
    ax.grid(True, alpha=0.3)
    
    # Rotate x-axis labels
    plt.xticks(rotation=45, ha='right')
    
    # Adjust every nth label to avoid crowding
    if len(dates) > 20:
        nth = len(dates) // 10
        for i, label in enumerate(ax.xaxis.get_ticklabels()):
            if i % nth != 0:
                label.set_visible(False)
    
    plt.tight_layout()
    
    # Save to BytesIO
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buffer.seek(0)
    
    return img_buffer


def get_index_name(mode: str) -> str:
    """Get the index name for a mode."""
    index_names = {
        'Hydrology': 'MNDWI',
        'Vegetation': 'NDRE',
        'WaterQuality': 'NDCI',
        'SoilVegetation': 'SAVI',
        'AlgaeBloom': 'FAI',
        'WaterRatio': 'WRI'
    }
    return index_names.get(mode, mode)


def get_mode_description(mode: str) -> str:
    """Get description for each analysis mode."""
    descriptions = {
        'Hydrology': 'Análisis de humedad y cuerpos de agua superficial mediante índice MNDWI',
        'Vegetation': 'Análisis de salud vegetativa mediante índice NDRE (clorofila)',
        'WaterQuality': 'Análisis de calidad de agua y turbidez mediante índice NDCI',
        'SoilVegetation': 'Análisis de vegetación ajustado por influencia del suelo (SAVI)',
        'AlgaeBloom': 'Detección de floraciones algales mediante índice FAI',
        'WaterRatio': 'Ratio agua-tierra mediante índice WRI'
    }
    return descriptions.get(mode, mode)


def generate_wetland_report(
    wetland_name: str,
    wetland_metadata: Dict[str, Any],
    analysis_results: Dict[str, Any],
    start_date: str,
    end_date: str
) -> io.BytesIO:
    """
    Generate a comprehensive Word report for wetland analysis.
    """
    # Create document
    doc = Document()
    
    # --- HEADER ---
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "WETLAND MONITOR - REPORTE DE ANÁLISIS"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.font.size = Pt(10)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(37, 99, 235)
    
    # --- TITLE ---
    title = doc.add_heading(f'Reporte de Análisis: {wetland_name}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- METADATA SECTION ---
    doc.add_heading('Información del Humedal', level=2)
    
    metadata_table = doc.add_table(rows=5, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    metadata_table.rows[0].cells[0].text = 'Nombre'
    metadata_table.rows[0].cells[1].text = wetland_name
    
    metadata_table.rows[1].cells[0].text = 'Región'
    metadata_table.rows[1].cells[1].text = wetland_metadata.get('region', 'N/A')
    
    metadata_table.rows[2].cells[0].text = 'Código'
    metadata_table.rows[2].cells[1].text = wetland_metadata.get('code', 'N/A')
    
    metadata_table.rows[3].cells[0].text = 'Coordenadas'
    metadata_table.rows[3].cells[1].text = wetland_metadata.get('coordinates', 'N/A')
    
    metadata_table.rows[4].cells[0].text = 'Fecha de Generación'
    metadata_table.rows[4].cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    doc.add_paragraph()
    
    # --- ANALYSIS PERIOD ---
    doc.add_heading('Período de Análisis', level=2)
    period_para = doc.add_paragraph()
    period_para.add_run(f'Desde: ').bold = True
    period_para.add_run(start_date)
    period_para.add_run(' | ')
    period_para.add_run(f'Hasta: ').bold = True
    period_para.add_run(end_date)
    
    doc.add_paragraph()
    
    # --- ANALYSIS RESULTS ---
    doc.add_heading('Resultados del Análisis', level=2)
    
    modes = ['Hydrology', 'Vegetation', 'WaterQuality', 'SoilVegetation', 'AlgaeBloom', 'WaterRatio']
    
    for mode in modes:
        if mode not in analysis_results:
            continue
            
        result = analysis_results[mode]
        stats = result.get('stats', {})
        maps = result.get('maps', {})
        
        # Mode heading
        doc.add_heading(f'{mode} - {get_index_name(mode)}', level=3)
        
        # Description
        doc.add_paragraph(get_mode_description(mode), style='Intense Quote')
        
        # Statistics table
        stats_table = doc.add_table(rows=7, cols=2)
        stats_table.style = 'Light List Accent 1'
        
        stats_table.rows[0].cells[0].text = 'Valor Actual (Mediana)'
        stats_table.rows[0].cells[1].text = f"{stats.get('current', 0):.4f}"
        
        stats_table.rows[1].cells[0].text = 'Valor Año Anterior'
        stats_table.rows[1].cells[1].text = f"{stats.get('last', 0):.4f}"
        
        trend = stats.get('trend', 0)
        trend_cell = stats_table.rows[2].cells[1]
        trend_text = f"{trend:+.2f}%"
        trend_cell.text = trend_text
        stats_table.rows[2].cells[0].text = 'Tendencia'
        
        # Color code trend
        if trend > 0:
            for run in trend_cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(34, 197, 94)  # Green
        else:
            for run in trend_cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(239, 68, 68)  # Red
        
        stats_table.rows[3].cells[0].text = 'Desviación Estándar'
        stats_table.rows[3].cells[1].text = f"{stats.get('current_std', 0):.4f}"
        
        stats_table.rows[4].cells[0].text = 'Coeficiente de Variación'
        stats_table.rows[4].cells[1].text = f"{stats.get('cv', 0):.2f}%"
        
        stats_table.rows[5].cells[0].text = 'Puntos de Datos'
        stats_table.rows[5].cells[1].text = str(stats.get('data_count', 0))
        
        stats_table.rows[6].cells[0].text = 'Valores Atípicos'
        stats_table.rows[6].cells[1].text = str(stats.get('outlier_count', 0))
        
        doc.add_paragraph()

        # --- MAPS & LEGEND ---
        doc.add_heading('Mapas del Índice', level=4)
        
        # Download images
        # Download images
        # Download images logic
        img_data_start = None
        img_data_end = None
        
        if 'start_year' in maps and 'thumb_url' in maps['start_year']:
            thumb_url_start = maps['start_year']['thumb_url']
            if thumb_url_start:
                img_data_start = download_image(thumb_url_start)

        if 'end_year' in maps and 'thumb_url' in maps['end_year']:
            thumb_url_end = maps['end_year']['thumb_url']
            if thumb_url_end:
                img_data_end = download_image(thumb_url_end)
        
        # Create map table (1 row, 2 columns) if at least one image exists
        if img_data_start or img_data_end:
            map_table = doc.add_table(rows=2, cols=2)
            map_table.autofit = True
            map_table.style = 'Table Grid' # Adds grid lines around everything
            
            # Row 0: Images
            cell_start = map_table.rows[0].cells[0]
            cell_end = map_table.rows[0].cells[1]
            
            # Row 1: Captions
            cell_cap_start = map_table.rows[1].cells[0]
            cell_cap_end = map_table.rows[1].cells[1]
            
            # Insert Start Map
            paragraph_start = cell_start.paragraphs[0]
            paragraph_start.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_data_start:
                run_start = paragraph_start.add_run()
                run_start.add_picture(img_data_start, width=Inches(2.8))
                # Add border via inline shape properties is complex in python-docx
                # A simple workaround for visual framing is using a table style or ensuring the picture has good contrast.
                # However, user requested "enmarcada". We can try adding a border to the picture simply by using a single-cell table with border for each,
                # but nested tables are tricky. 
                # Let's rely on the clean spacing for now or add a simple border run property if feasible.
                # Actually, standard python-docx doesn't support easy image borders.
                # We will use the 'Table Grid' style for the layout table itself if we want frames around cells.
                
                cell_cap_start.text = f"Mapa Inicial ({start_date})"
                cell_cap_start.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Insert End Map
            paragraph_end = cell_end.paragraphs[0]
            paragraph_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_data_end:
                run_end = paragraph_end.add_run()
                run_end.add_picture(img_data_end, width=Inches(2.8))
                
                cell_cap_end.text = f"Mapa Final ({end_date})"
                cell_cap_end.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Legend (Shared for both)
        legend_data = create_legend_image(mode)
        if legend_data:
            doc.add_picture(legend_data, width=Inches(5))
            doc.add_paragraph("Escala de Valores (Válida para ambos mapas)", style='Caption')

        
        # Add temporal chart
        doc.add_heading('Evolución Temporal', level=4)
        time_series = result.get('time_series', [])
        if time_series:
            chart_buffer = create_temporal_chart(time_series, mode)
            if chart_buffer:
                doc.add_picture(chart_buffer, width=Inches(6))
        
        doc.add_page_break()
    
    # --- FOOTER ---
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generado por WETLAND MONITOR | {datetime.now().strftime('%Y-%m-%d')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer
